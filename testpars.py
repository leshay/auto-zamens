import requests,time,os
from bs4 import BeautifulSoup
import docx, yaml, json, re, argparse
import main

def task_update():
    url = 'http://www.gtec-bks.by/'
    response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})

    if response.status_code == 200:
        html = response.text

        # Используйте BeautifulSoup для анализа HTML
        soup = BeautifulSoup(html, 'html.parser')

        # Найдите нужный кусок кода, заменив 'tag' и 'class' на соответствующие значения
        target_element = (soup.find('div', class_='col-md-9').find('div', class_='mobile-side-buttons')
                          .find('a'))
        saved_html = open('output.html', 'r', encoding='utf-8').read()
        if saved_html == target_element.prettify():
            print('замен нету')
        else:
            # Получаем список файлов в папке temp
            files = os.listdir('temp')
            # Проверяем, есть ли файлы в папке
            if not files:
                handler_file(download_pdf(f'http://www.gtec-bks.by{target_element["href"]}', target_element.text[-10:]))
                with open('output.html', 'w', encoding='utf-8') as save_element:
                    save_element.write(target_element.prettify())
            else:
                local_time = time.ctime(seconds)
                print('В папки существуют файлы такого не должно быть', local_time)

def download_pdf(url, destination):
    import urllib.request
    urllib.request.urlretrieve(f"{url}", f"temp/{destination}.doc")

    return f'temp/{destination}.doc'

def handler_file(filename, convert=True):
    converter_path = conver(filename) if convert else filename
    start_time = time.time()
    zamen = read_docx(converter_path)
    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f'Время выполнение zamen: {elapsed_time}')
    print(zamen)
    start_time_new = time.time()
    zamen_test = test_dock_reader(converter_path)
    end_time_new = time.time()
    elapsed_time_new = end_time_new - start_time_new
    print(f'Время выполнение zamen_test: {elapsed_time_new}')
    print(zamen_test)
    start_time_new_mashed = time.time()
    zamen_test_mashed = test_dock_mashd_reader(converter_path)
    end_time_new_mashed = time.time()
    elapsed_time_new_mashed = end_time_new_mashed - start_time_new_mashed
    print(f'Время выполнение zamen_test_mashed: {elapsed_time_new_mashed}')
    print(zamen_test_mashed)

    if convert is False:
        pass
    else:
        # Получаем все файлы в папке
        files = os.listdir('temp')

        # Перебираем файлы и удаляем каждый
        for file_name in files:
            file_path = os.path.join('temp', file_name)
            if os.path.isfile(file_path):
                os.remove(file_path)
    main.send_file_back_to_previous()



def conver(filename):
    import convertio
    import time
    try:
        client = convertio.Client(token="036c8042eca7be91df521e84f67836ee")
        conversion_id = client.convert_by_filename(f'{filename}', 'docx')
        while client.check_conversion(conversion_id).step != 'finish':
            time.sleep(1)
        client.download(conversion_id, f'{filename}x')
        client.delete(conversion_id)
        return f'{filename}x'
    except:
        print('Внимание ошибка с основным токеном')
        print('Использую дополнительные токены')
        client = convertio.Client(token="de57dba61142b5a7373fff6e2873e67d")
        conversion_id = client.convert_by_filename(f'{filename}', 'docx')
        while client.check_conversion(conversion_id).step != 'finish':
            time.sleep(1)
        client.download(conversion_id, f'{filename}x')
        client.delete(conversion_id)
        return f'{filename}x'

def read_docx(filename):
    doc = docx.Document(f'{filename}')
    data = {}
    data['data'] = filename[5:15]
    # Читаем данные из таблицы
    for table in doc.tables:
        try:
            # читаем данные из таблицы
            for row in table.rows[1:]:
                row_data = []

                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    cleaned_row = [value if not isinstance(value, str) else ' '.join(word.strip() for word in value.split()) for
                                   value in row_data]

                    to_append = cleaned_row[2:] if len(cleaned_row) > 3 else cleaned_row[2]
                    if cleaned_row[0] not in data:
                        data[cleaned_row[0]] = {cleaned_row[1]: to_append}
                    else:
                        data[cleaned_row[0]][cleaned_row[1]] = to_append
        except Exception:
            pass

    with open("version_output/output.yaml", "w", encoding='UTF-8') as file:
        yaml.dump(data, file, allow_unicode=True, default_flow_style=False)
    with open("version_output/output.json", "w", encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)
    return True

def test_dock_reader(filename):
    doc = docx.Document(f'{filename}')
    data = {}
    data['data'] = filename[5:15]
    # Читаем данные из таблицы
    for table in doc.tables:
        try:
            # Читаем данные из таблицы
            for row in table.rows[1:]:
                row_data = []
                for cell in [row.cells[0], row.cells[1], row.cells[3]]:
                    cell_text = cell.text.strip()
                    if not cell_text:  # Если значение пустое
                        break  # Прерываем цикл
                    if '\n' in cell_text:
                        cell_text = cell_text.split('\n')
                        row_data.extend(cell_text)
                    else:
                        row_data.append(cell_text)

                if len(row_data) > 2 and 'Н/Б' in row_data[2]:
                    row_data[2] = row_data[2].replace('Н/Б', 'None')
                    row_data[2] = ' '.join(row_data[2].split())

                if len(row_data) > 3 and 'Н/Б' in row_data[3]:
                    row_data[3] = row_data[3].replace('Н/Б', 'None')
                    row_data[3] = ' '.join(row_data[3].split())

                for index, element in enumerate(row_data):
                    if isinstance(element, str) and " " in element:
                        if element.find('п/гр') != -1:
                            row_data[index - 1] += ' (п/гр)'
                        if "(" not in element and ")" not in element:  # Добавляем проверку на наличие ()
                            split_element = element.split("   ")[0]  # Удаляем все значения после первого пробела
                            if not split_element.strip():
                                del row_data[index]
                            else:
                                row_data[index] = split_element
                if row_data:
                    row_data[2] = re.sub(
                        r'^\d+/\d+\b|\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.|[А-ЯЁ][а-яё]+\s*[А-ЯЁ]\.(?:\s*[\w"]|\.\s*[\w"])',
                        '', row_data[2])
                    if len(row_data) > 3:
                        row_data[3] = re.sub(
                            r'^\d+/\d+\b|\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.|[А-ЯЁ][а-яё]+\s*[А-ЯЁ]\.(?:\s*[\w"]|\.\s*[\w"])',
                            '', row_data[3])

                    if len(row_data) > 3:
                        if re.match(r'^\([^()]+\)$', row_data[3].strip()):
                            row_data[2] = row_data[2] + row_data[3]
                            del row_data[3]
                        else:
                            pass

                    for index, value in enumerate(row_data):
                        if index == 2 or index == 3:  # Проверяем индексы 2 и 3
                            row_data[index] = value.replace('Мирошниченко И', '').replace('Васьковцова Н.', '').replace(
                                'Железнякова И.', '').replace('Евдокимович В.', '').replace('Вальченко', '').replace(
                                'Кузьменков А.', '').replace('Мальдис Н.', '').replace('Потапнева Я.', '')

                    cleaned_row = [
                        ' '.join(word.strip() for word in value.split()) if isinstance(value, str) else value
                        for value in row_data if value != '.' and value.strip()  # Проверка на пустой элемент
                    ]

                    to_append = cleaned_row[2:] if len(cleaned_row) > 3 else cleaned_row[2]
                    if cleaned_row[0] not in data:
                        if isinstance(cleaned_row[1], (int, float)) or cleaned_row[1].lstrip('-').replace('-', '').replace('.', '',1).isdigit():
                            data[cleaned_row[0]] = {cleaned_row[1]: to_append}
                        else:
                            data[cleaned_row[0]] = to_append
                    else:
                        if isinstance(cleaned_row[1], (int, float)) or cleaned_row[1].lstrip('-').replace('-', '').replace('.', '',1).isdigit():
                            data[cleaned_row[0]][cleaned_row[1]] = to_append
                        else:
                            data[cleaned_row[0]] = to_append

        except Exception:
            pass
    with open("version_output/output_result.yaml", "w", encoding='UTF-8') as file:
        yaml.dump(data, file, allow_unicode=True, default_flow_style=False)
    with open("version_output/output_result.json", "w", encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)
    return True

def test_dock_mashd_reader(filename):
    import mashed_test
    doc = docx.Document(f'{filename}')
    data = {}
    data['data'] = filename[5:15]
    cab = {}
    cabinets = test_wort(filename)
    if cabinets:
        for cabinet in cabinets:
            cab.setdefault(cabinet[0], {})[cabinet[1]] = int(cabinet[2]) if cabinet[2].isdigit() else cabinet[2]
        data['cabinets'] = cab

    # Читаем данные из таблицы
    for table in doc.tables:
        try:
            for row in table.rows[1:]:
                row_data = []
                for idx, cell in enumerate([row.cells[0], row.cells[1], row.cells[3], row.cells[4]]):
                    # Если это row.cells[4], обрабатываем без .strip()
                    if idx == 3:  # row.cells[4] — это четвертый элемент в списке
                        cell_text = cell.text  # Без .strip()
                    else:
                        cell_text = cell.text.strip()  # С .strip() для остальных

                    if not cell_text:  # Если значение пустое
                        pass
                    elif '\n' in cell_text:
                        cell_text = cell_text.split('\n')
                        row_data.append(cell_text)
                    else:
                        row_data.append(cell_text)

                if len(row_data) <= 2:
                    pass
                else:
                    if type(row_data[2]) == list:
                        row_data[2] = [item.replace('Н/Б', 'None').strip() if 'Н/Б' in item else item for item in row_data[2]]
                    else:
                        if 'Н/Б' in row_data[2]:
                            row_data[2] = row_data[2].replace('Н/Б', 'None')
                            row_data[2] = ' '.join(row_data[2].split())

                    # Проверка и обработка row_data[2]
                    if isinstance(row_data[2], list):
                        # Если row_data[2] — список, проходим по его элементам
                        row_data[2] = [
                            element.split("   ")[0] if isinstance(element,
                                                                  str) and " " in element and "(" not in element and ")" not in element
                            else element for element in row_data[2]
                        ]
                        # Убираем лишние пробелы в каждом элементе списка перед обработкой
                        row_data[2] = [element.strip() if isinstance(element, str) else element for element in
                                       row_data[2]]
                        # Логика для объединения второго элемента, если он в скобках, с первым
                        if len(row_data[2]) > 1 and isinstance(row_data[2][1], str) and row_data[2][1].startswith(
                                '(') and \
                                row_data[2][1].endswith(')'):
                            # Убираем лишние пробелы перед объединением
                            row_data[2][0] += ' ' + row_data[2][1]  # Добавляем второй элемент к первому
                            row_data[2] = row_data[2][
                                0]  # Заменяем список на строку, состоящую из первого и второго элемента

                        # Дополнительная логика для 'п/гр' внутри списка
                        for idx, element in enumerate(row_data[2]):
                            if isinstance(element, str) and 'п/гр' in element and idx > 0:
                                row_data[2][idx - 1] += ' (п/гр)'

                    elif isinstance(row_data[2], str) and " " in row_data[2]:
                        # Если row_data[2] — строка, применяем ту же логику
                        if 'п/гр' in row_data[2]:
                            row_data[1] += ' (п/гр)'  # Добавляем к предыдущему элементу
                        if "(" not in row_data[2] and ")" not in row_data[2]:
                            split_element = row_data[2].split("   ")[0]
                            if not split_element.strip():
                                row_data[2] = ''  # Заменяем на пустую строку, если элемент пустой
                            else:
                                row_data[2] = split_element



                    if len(row_data) == 4:
                        if row_data[3] is not None:
                            if type(row_data[2]) is not type(row_data[3]):
                                if type(row_data[3]) is list:
                                    # Удаляем пустые элементы из списка
                                    row_data[3] = [item for item in row_data[3] if item]

                                    # Если в списке остался только один элемент, преобразуем его в строку
                                    if len(row_data[3]) == 1:
                                        row_data[3] = row_data[3][0]
                        if isinstance(row_data[3], list):
                            row_data[3] = [item if item else '-' for item in row_data[3]]
                        else:
                            row_data[3] = row_data[3] if row_data[3] else '-'



                    if type(row_data[2]) == list:
                        row_data[2] = [re.sub(
                            r'^\d+/\d+\b|\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.|[А-ЯЁ][а-яё]+\s*[А-ЯЁ]\.(?:\s*[\w"]|\.\s*[\w"])',
                            '', item) for item in row_data[2]]
                    else:
                        row_data[2] = re.sub(
                            r'^\d+/\d+\b|\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.|[А-ЯЁ][а-яё]+\s*[А-ЯЁ]\.(?:\s*[\w"]|\.\s*[\w"])',
                            '', row_data[2])

                    # Теперь проверяем только элементы внутри row_data[2] для замены строк
                    if type(row_data[2]) == list:
                        row_data[2] = [item.replace('Мирошниченко И', '').replace('Васьковцова Н.', '').replace(
                            'Железнякова И.', '').replace('Евдокимович В.', '').replace('Вальченко', '').replace(
                            'Кузьменков А.', '').replace('Мальдис Н.', '').replace('Потапнева Я.', '') for item in
                                       row_data[2]]
                    else:
                        row_data[2] = row_data[2].replace('Мирошниченко И', '').replace('Васьковцова Н.', '').replace(
                            'Железнякова И.', '').replace('Евдокимович В.', '').replace('Вальченко', '').replace(
                            'Кузьменков А.', '').replace('Мальдис Н.', '').replace('Потапнева Я.', '')

                    cleaned_row = [
                        ' '.join(word.strip() for word in value.split()) if isinstance(value,
                                                                                       str) and value.strip() and value not in [
                                                                                '.', '.,'] else (
                            [
                                ' '.join(word.strip() for word in item.split()) if isinstance(item,
                                                                                              str) and item.strip() and item not in [
                                                                                       '.', '.,'] else item
                                for item in value if
                                item and isinstance(item, str) and item.strip() and item not in ['.', '.,']
                            ][0] if isinstance(value, list) and len([
                                ' '.join(word.strip() for word in item.split()) if isinstance(item,
                                                                                              str) and item.strip() and item not in [
                                                                                       '.', '.,'] else item
                                for item in value if
                                item and isinstance(item, str) and item.strip() and item not in ['.', '.,']
                            ]) == 1 else [
                                ' '.join(word.strip() for word in item.split()) if isinstance(item,
                                                                                              str) and item.strip() and item not in [
                                                                                       '.', '.,'] else item
                                for item in value if
                                item and isinstance(item, str) and item.strip() and item not in ['.', '.,']
                            ]
                        )
                        for value in row_data if value and (
                                (isinstance(value, str) and value.strip() and value not in ['.', '.,']) or isinstance(
                            value, list)
                        )
                    ]
                    if len(cleaned_row) == 4:
                        if cleaned_row[3] is not None:
                            if isinstance(cleaned_row[3], list) and not cleaned_row[3]:
                                cleaned_row.pop(3)  # Удаляет элемент с индексом 3

                    if isinstance(cleaned_row[2], str):
                        if re.search(r"\b1\s*ч\b", cleaned_row[2]):
                            cleaned_row[2] = [cleaned_row[2], "None"]
                            if isinstance(cleaned_row[3], str):
                                cleaned_row[3] = [cleaned_row[3], "-"]
                    if cleaned_row:
                        if type(cleaned_row[2]) == list:
                            cleaned_row[2] = [mashed_test.testing_obl(cleaned_row) for cleaned_row in cleaned_row[2]]
                        else:
                            cleaned_row[2] = mashed_test.testing_obl(cleaned_row[2])

                        if len(cleaned_row) == 4:
                            if isinstance(cleaned_row[3], list):  # для английского
                                if cleaned_row[2] == 'Английский язык':
                                    cleaned_row[3] = '/'.join([item for item in cleaned_row[3]])
                        try:
                            to_fuxing = cleaned_row[4:6] if len(cleaned_row) > 5 else cleaned_row[3]
                        except IndexError:
                            to_fuxing = ''

                        to_append = cleaned_row[2:4] if len(cleaned_row) > 5 else cleaned_row[2]
                        if to_fuxing != '':
                            formatted_value = f'{to_append} [{to_fuxing}]'
                            if type(to_append) == list:
                                formatted_value = [f"{x} [{y}]" for x, y in zip(to_append, to_fuxing)]
                        else:
                            formatted_value = to_append

                        if cleaned_row[0] not in data:
                            if isinstance(cleaned_row[1], (int, float)) or cleaned_row[1].lstrip('-').replace('-', '').replace('.', '',1).isdigit():
                                data[cleaned_row[0]] = {cleaned_row[1]: formatted_value}
                            else:
                                data[cleaned_row[0]] = formatted_value
                        else:
                            if isinstance(cleaned_row[1], (int, float)) or cleaned_row[1].lstrip('-').replace('-', '').replace('.', '',1).isdigit():
                                data[cleaned_row[0]][cleaned_row[1]] = formatted_value
                            else:
                                data[cleaned_row[0]] = formatted_value

        except Exception:
            pass
    with open("version_output/output_result_mashed.yaml", "w", encoding='UTF-8') as file:
        yaml.dump(data, file, allow_unicode=True, default_flow_style=False, sort_keys=False)
    with open("version_output/output_result_mashed.json", "w", encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)
    return True

def test_wort(filename):
    doc = docx.Document(f'{filename}')
    # Исходный текст
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    # Объединяем текст в одну строку
    full_text = '\n'.join(text)
    # Регулярное выражение для поиска группы "Б-21 (4) 221"
    pattern = r'\b([А-Я]-\d{1,2})\s*\(([\d\\]+)\)\s*(\d{3})\b'
    # Ищем все совпадения
    matches = re.findall(pattern, full_text)
    # Преобразуем данные, заменяя \\ на / во втором элементе
    matches = [(x[0], x[1].replace('\\', '/'), x[2]) for x in matches]
    return matches

def tubler(arguments = None):
    if arguments:
        if arguments == 'now':
            # Получаем все файлы в папке
            files = os.listdir('temp')

            # Фильтруем только файлы с расширением .docx
            docx_files = [file for file in files if file.endswith('.docx')]

            # Выводим файлы и проверяем их количество
            if len(docx_files) > 1:
                print(f"Найдено несколько файлов: {docx_files}")
            elif len(docx_files) == 1:
                print(f"Найден файл: {docx_files[0]}")
                yn = input('Продолжаем? y/n: ').strip().lower()
                if yn == 'y':
                    handler_file(filename=f'temp/{docx_files[0]}', convert=False)
                elif yn == 'n':
                    pass
                else:
                    print("Некорректный ввод. Пожалуйста, введите 'y' или 'n'.")
            else:
                print("Файлы .docx не найдены.")
        else:
            handler_file(filename=arguments,convert=False)
    else:
        task_update()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Ваши варианты')
    parser.add_argument('-f', '--file', type=str, help='Имя файла или либо из сохранения используя now', default=None)
    args = parser.parse_args()

    tubler(args.file)